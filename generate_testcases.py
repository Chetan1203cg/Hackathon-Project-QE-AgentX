"""
Testcase Generator
==================
Reads a Jira RSS XML (user story) and generates a Word document
containing structured test cases, following the standard template.

Usage:
    python generate_testcases.py <path-to-jira.xml> [--output <output.docx>]
    python generate_testcases.py <JIRA-KEY> [--output <output.docx>]

Requirements:
    pip install python-docx lxml requests
"""

import argparse
import json
import os
import re
import sys
from html.parser import HTMLParser
from pathlib import Path

from lxml import etree
from docx import Document


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

class _HTMLTextExtractor(HTMLParser):
    """Strip HTML tags and return plain text."""
    def __init__(self):
        super().__init__()
        self._parts = []

    def handle_data(self, data):
        self._parts.append(data)

    def handle_entityref(self, name):
        pass

    def get_text(self):
        return " ".join(self._parts).strip()


def strip_html(html: str) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(html)
    return re.sub(r"\s+", " ", parser.get_text()).strip()


# ---------------------------------------------------------------------------
# XML parsing
# ---------------------------------------------------------------------------

def parse_jira_xml(xml_path: str) -> dict:
    """Extract relevant fields from a Jira RSS XML export."""
    tree = etree.parse(xml_path)
    root = tree.getroot()

    item = root.find(".//item")
    if item is None:
        raise ValueError("No <item> element found in XML")

    def text(tag):
        el = item.find(tag)
        return (el.text or "").strip() if el is not None else ""

    key = text("key")
    summary = text("summary")
    link = text("link")
    issue_type = text("type")
    priority = text("priority")
    status = text("status")
    component = text("component")
    assignee = text("assignee")

    raw_description = text("description")
    description = strip_html(raw_description)

    return {
        "key": key,
        "summary": summary,
        "link": link,
        "type": issue_type,
        "priority": priority,
        "status": status,
        "component": component,
        "assignee": assignee,
        "description": description,
    }


# ---------------------------------------------------------------------------
# LLM call
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """You are an expert QA engineer. Your task is to analyze a Jira user story
and generate thorough, structured test cases.

IMPORTANT: Always respond in English, regardless of the language of the input story.

Return ONLY valid JSON — no markdown, no extra text.
The JSON must follow this exact schema:

{
  "feature": "<short feature name, max 60 chars>",
  "test_scenario": "<one sentence describing the overall test scenario>",
  "scenario_description": "<one sentence summary of what is being tested>",
  "acceptance_criteria": ["<AC item 1>", "<AC item 2>", ...],
  "test_cases": [
    {
      "title": "<Test Case N – Short Title>",
      "description": "<what this test case validates>",
      "preconditions": ["<precondition 1>", ...],
      "steps": ["<step 1>", "<step 2>", ...],
      "expected_result": "<clear expected outcome>"
    }
  ]
}

Rules:
- Create one test case per acceptance criterion, plus edge cases where appropriate.
- Test case titles must be SHORT precise headlines (max 6 words after "Test Case N –"), e.g. "Test Case 1 – Cart Flyout Display".
  Do NOT paste the full AC text into the title.
- Preconditions must be concrete and testable (e.g. "User is logged in as Admin").
- Each step must be a clear, concise USER ACTION (e.g. "Click the Submit button", "Navigate to Settings page").
  Do NOT repeat the acceptance criteria text in the steps.
  Do NOT write steps like "Validate AC: <ac text>" — instead describe what the tester physically does.
  Verification steps should start with "Verify that …" and state the observable outcome.
- Steps must be numbered actions (do NOT include numbers in the text — they will be auto-numbered).
- Expected results must be clear, specific, and verifiable.
- Use plain language. No HTML.
"""


def build_prompt_for_copilot(story: dict) -> str:
    """Create a ready-to-paste prompt for Copilot chat or any LLM UI."""
    system_context = load_system_context(Path(__file__).parent, story.get('component', ''))
    prompt_parts = [SYSTEM_PROMPT]
    if system_context:
        prompt_parts.append(system_context)
    prompt_parts += [
        "",
        "Analyze the following Jira user story and return ONLY JSON.",
        "",
        f"Jira Issue: {story['key']}",
        f"Summary: {story['summary']}",
        f"Component: {story['component']}",
        f"Priority: {story['priority']}",
        f"Status: {story['status']}",
        f"Link: {story['link']}",
        "",
        "Description / User Story & Acceptance Criteria:",
        story['description'],
    ]
    return "\n".join(prompt_parts)


# ---------------------------------------------------------------------------
# Word document generation
# ---------------------------------------------------------------------------

def add_heading(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, style: str = "Normal"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc: Document, text: str, index: int):
    # Use explicit numbering so each test case can restart from 1.
    doc.add_paragraph(f"{index}. {text}", style="Normal")


def build_document(story: dict, test_data: dict, output_path: str):
    """Create a .docx following the standard test case template."""
    doc = Document()

    # --- Document header ---
    add_heading(doc, f"Test Cases – {test_data['feature']}", level=1)
    add_heading(doc, test_data["feature"], level=2)
    add_heading(doc, f"User story: {story['key']} {story['summary']}", level=2)
    add_heading(doc, f"URL: {story['link']}", level=2)

    doc.add_paragraph()

    add_paragraph(doc, f"Test scenario: {test_data['test_scenario']}")
    add_paragraph(doc, f"Description: {test_data['scenario_description']}")
    add_paragraph(doc, "AC:")
    for ac in test_data.get("acceptance_criteria", []):
        add_bullet(doc, ac)

    # --- Individual test cases ---
    for tc in test_data.get("test_cases", []):
        add_heading(doc, tc["title"], level=3)
        add_paragraph(doc, f"Description: {tc['description']}")
        add_paragraph(doc, "Preconditions:")
        for pre in tc.get("preconditions", []):
            add_bullet(doc, pre)
        add_paragraph(doc, "Test Steps:")
        for step_index, step in enumerate(tc.get("steps", []), start=1):
            add_numbered(doc, step, step_index)
        add_paragraph(doc, f"Expected Result: {tc['expected_result']}")
        add_paragraph(doc, "Test Results:")
        add_paragraph(doc, "[expected text and screenshots]")

    doc.save(output_path)
    print(f"Document saved: {output_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def load_config(config_path: str) -> dict:
    with open(config_path, "r", encoding="utf-8") as f:
        cfg = json.load(f)
    jira = cfg.setdefault("jira", {})
    if not jira.get("api_token"):
        jira["api_token"] = os.environ.get("JIRA_API_TOKEN", "")
    return cfg


def load_test_data_json(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


_CONTEXT_ROUTING: list = [
    # (list of component keywords to match, context filename)
    (["mydealer", "my dealer", "my-dealer"], "system_context.json"),
    (["cms", "evhub", "ev hub", "ev-hub", "onehub", "one hub"], "system_context_cms.json"),
]


def _pick_context_file(component: str) -> str:
    """Return the best-matching context filename for the given component string."""
    comp_lower = (component or "").lower()
    for keywords, filename in _CONTEXT_ROUTING:
        if any(kw in comp_lower for kw in keywords):
            return filename
    return "system_context.json"  # default fallback


def load_system_context(script_dir: Path, component: str = "") -> str:
    """Load system context from data/ and format it for the LLM prompt.

    The context file is chosen automatically based on the Jira component field.
    """
    filename = _pick_context_file(component)
    ctx_path = script_dir / "data" / filename
    if not ctx_path.exists():
        # Try the default fallback
        ctx_path = script_dir / "data" / "system_context.json"
    if not ctx_path.exists():
        return ""
    with open(ctx_path, "r", encoding="utf-8") as f:
        ctx = json.load(f)
    lines = [
        f"## System Context: {ctx.get('system', 'Unknown')}",
        f"Documentation: {ctx.get('confluence_doc', '')}",
        "",
        ctx.get("summary", ""),
        "",
        "### Key Concepts",
    ]
    for key, value in ctx.get("key_concepts", {}).items():
        lines.append(f"- **{key.replace('_', ' ').title()}**: {value}")
    preconditions = ctx.get("common_preconditions", [])
    if preconditions:
        lines += ["", "### Typical Preconditions"]
        for pre in preconditions:
            lines.append(f"- {pre}")
    return "\n".join(lines)


_JIRA_KEY_RE = re.compile(r'^[A-Z][A-Z0-9]+-\d+$')
_JIRA_BASE_URL = "https://vwgroup-b2b.atlassian.net"


def main():
    parser = argparse.ArgumentParser(
        description="Generate test-case Word document from a Jira RSS XML user story or Jira issue key."
    )
    parser.add_argument("xml", help="Path to the Jira RSS XML file, or a Jira issue key (e.g. NGWD6-46920)")
    parser.add_argument(
        "--output", "-o",
        help="Output .docx path (default: output/<ISSUE_KEY>.docx)",
    )
    parser.add_argument(
        "--config", "-c",
        default=str(Path(__file__).parent / "config.json"),
        help="Path to config.json (default: config.json next to this script)",
    )
    parser.add_argument(
        "--emit-prompt",
        help=(
            "Write a Copilot/LLM prompt text file and exit (no API key needed). "
            "Example: --emit-prompt output/prompt.txt"
        ),
    )
    parser.add_argument(
        "--from-json",
        help="Use pre-generated test-case JSON instead of calling the pipeline.",
    )
    args = parser.parse_args()

    # --- Jira issue key shortcut ---
    if _JIRA_KEY_RE.match(args.xml):
        import jira_to_testcases as jtt  # lazy import avoids circular dependency

        jira_link = f"{_JIRA_BASE_URL}/browse/{args.xml}"
        if not os.path.exists(args.config):
            print(f"ERROR: Config file not found: {args.config}", file=sys.stderr)
            sys.exit(1)
        config = load_config(args.config)

        print(f"Fetching Jira issue: {args.xml}")
        story = jtt.fetch_jira_story(jira_link, config)
        print(f"  Issue  : {story['key']} – {story['summary']}")

        if args.output:
            output_path = args.output
        else:
            output_dir = Path(__file__).parent / "output"
            output_dir.mkdir(exist_ok=True)
            output_path = str(output_dir / f"{story['key']}.docx")

        print("Generating test cases...")
        test_data = jtt.rule_based_test_data(story)
        print("  Used rule-based generation")
        print(f"  Generated {len(test_data.get('test_cases', []))} test case(s)")
        build_document(story, test_data, output_path)
        return

    # --- XML file mode ---
    print(f"Parsing: {args.xml}")
    story = parse_jira_xml(args.xml)
    print(f"  Issue  : {story['key']} – {story['summary']}")

    # Determine output path
    if args.output:
        output_path = args.output
    else:
        output_dir = Path(__file__).parent / "output"
        output_dir.mkdir(exist_ok=True)
        output_path = str(output_dir / f"{story['key']}.docx")

    # Copilot-assisted prompt export mode
    if args.emit_prompt:
        prompt = build_prompt_for_copilot(story)
        prompt_path = Path(args.emit_prompt)
        prompt_path.parent.mkdir(parents=True, exist_ok=True)
        prompt_path.write_text(prompt, encoding="utf-8")
        print(f"Prompt file saved: {prompt_path}")
        print("Next: paste prompt into Copilot chat, save JSON output, then run with --from-json")
        return

    # JSON input mode
    if args.from_json:
        print(f"Loading test data JSON: {args.from_json}")
        test_data = load_test_data_json(args.from_json)
        build_document(story, test_data, output_path)
        return

    if not os.path.exists(args.config):
        print(f"ERROR: Config file not found: {args.config}", file=sys.stderr)
        sys.exit(1)
    config = load_config(args.config)

    import jira_to_testcases as jtt2
    test_data = jtt2.rule_based_test_data(story)
    print(f"  {len(test_data.get('test_cases', []))} test case(s) generated")

    build_document(story, test_data, output_path)


if __name__ == "__main__":
    main()
